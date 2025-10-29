import docx
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from io import BytesIO
from docx.shared import Pt
import logging


def turn_changes_into_word(changes: dict, tour_id_to_df: dict, optimized_distances: dict) -> str:
    """
    Create word document from changes dict.
    Each tour gets its own page.

    Args:
        changes: Dictionary with tour_id as key and list of change descriptions as value
        tour_id_to_df: Dictionary with tour_id as key and dict with keys:
                       - "tour_df": pd.DataFrame with columns children_on_tour, Street, Number, PLZ, Region
                       - "symbol": tour symbol
                       - "km_besetzt": distance in km
        optimized_distances: dict mapping tour_id to optimized distance
    """
    # Neues Dokument
    doc = docx.Document()

    # Querformat einstellen
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    # Überschrift
    doc.add_heading("Tourenänderung zur Streckenoptimierung", 0)

    for tour_id, change_list in changes.items():
        tour_data = tour_id_to_df[tour_id]
        doc.add_heading(f"Änderungen an Tour {tour_id[-5:]} - {tour_data['symbol']}", level=1)

        if not change_list:
            doc.add_paragraph("Keine Änderungen für diese Tour.")
        else:
            doc.add_paragraph(change_list)

        # Füge optimierte Distanz hinzu
        doc.add_paragraph(f"Original Malteser Distanz: {tour_data['km_besetzt']} km")
        if tour_id in optimized_distances:
            doc.add_paragraph(f"Optimierte Distanz: {optimized_distances[tour_id]} km")

        doc.add_page_break()

    # Save file in acceptable format
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

def turn_df_into_word(tour_id_to_df: dict, google_distances=None, optimized_distances=None) -> str:
    """
    Create word document from dataframe.
    5 Tours per Page.

    Args:
        tour_id_to_df: Dictionary with tour_id as key and dict with keys:
                       - "tour_df": pd.DataFrame with columns children_on_tour, Street, Number, PLZ, Region
                       - "symbol": tour symbol
                       - "km_besetzt": distance in km
        filename: output filename
        google_distances: dict mapping tour_id to google maps distance
        optimized_distances: dict mapping tour_id to optimized distance
    """
    if google_distances is None:
        google_distances = {}
    if optimized_distances is None:
        optimized_distances = {}

    # Neues Dokument
    doc = docx.Document()

    # Querformat einstellen
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    # Überschrift
    original_distances_all = sum([tour["km_besetzt"] for tour in tour_id_to_df.values()])
    doc.add_heading("Tourenübersicht der Maria Stern Schule", 0)
    doc.add_heading(f"Gesamtdistanz Malteser: {original_distances_all} km", level=1)
    if len(google_distances) != 0:
        doc.add_heading(f"Gesamtdistanz OSM {sum(google_distances.values())} km", level=2)
    if len(optimized_distances) != 0:
        doc.add_heading(f"Gesamtdistanz Optimiert: {sum(optimized_distances.values())} km", level=2)

    max_tours_per_page = 5
    tour_ids = list(tour_id_to_df.keys())
    num_tours = len(tour_ids)

    # In Seitenblöcke teilen
    for start in range(0, num_tours, max_tours_per_page):
        end = min(start + max_tours_per_page, num_tours)
        tours_block = tour_ids[start:end]

        num_rows = 9 + 1 + (4 if len(google_distances) != 0 else 2)  # 9 Kinder/Zeilen, 1 Anm., 2 Zusatzzeilen
        if len(optimized_distances) != 0:
            num_rows += 2
        num_cols = len(tours_block) + 1  # +1 für erste Spalte (Labels)

        table = doc.add_table(rows=num_rows + 1, cols=num_cols)  # +1 für Kopfzeile
        table.style = "Table Grid"
        table.autofit = False

        # Define Header with enumeration (continues across pages)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ""
        for col_idx, tour_id in enumerate(tours_block, start=start + 1):
            hdr_cells[col_idx - start].text = str(col_idx)
            p = hdr_cells[col_idx - start].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True

        # First row with Tour informations
        fr_cells = table.rows[1].cells
        fr_cells[0].text = "Tour-Infos:"
        for i, tour_id in enumerate(tours_block, start=1):
            tour_data = tour_id_to_df[tour_id]
            fr_text = f"{tour_id[-5:]} - {tour_data['symbol']}"  # Letzte Ziffern der ID + Symbol
            fr_cells[i].text = fr_text
            p = fr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True

        # Erste Spalte: Labels
        for i in range(1, 10):
            table.rows[i + 1].cells[0].text = str(i)
        table.rows[10].cells[0].text = "Anm."

        # KM-Besetzt Zeile
        row_km = table.rows[11].cells
        merged_km = row_km[0].merge(row_km[-1])
        p = merged_km.paragraphs[0]
        p.add_run("KM-Besetzt Malteser").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Grüne Zeile für Malteser
        row_green_malteser = table.rows[12].cells

        # Dynamische Zeilenzuweisung für OSM und Optimiert
        if len(google_distances) != 0:
            # OSM nutzt Zeilen 13-14
            row_km_osm = table.rows[13].cells
            merged_km_osm = row_km_osm[0].merge(row_km_osm[-1])
            p = merged_km_osm.paragraphs[0]
            p.add_run("KM-Besetzt OSM").bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_green_osm = table.rows[14].cells

            if len(optimized_distances) != 0:
                # Optimiert nutzt Zeilen 15-16
                row_km_opt = table.rows[15].cells
                merged_km_opt = row_km_opt[0].merge(row_km_opt[-1])
                p = merged_km_opt.paragraphs[0]
                p.add_run("KM-Besetzt optimiert").bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_green_optimized = table.rows[16].cells
            else:
                row_green_optimized = None
        else:
            # Wenn keine OSM-Daten vorhanden, nutzt Optimiert Zeilen 13-14
            row_green_osm = None
            if len(optimized_distances) != 0:
                row_km_opt = table.rows[13].cells
                merged_km_opt = row_km_opt[0].merge(row_km_opt[-1])
                p = merged_km_opt.paragraphs[0]
                p.add_run("KM-Besetzt optimiert").bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_green_optimized = table.rows[14].cells
            else:
                row_green_optimized = None

        # Grüne Zeilen einfärben
        for row in [row for row in [row_green_malteser, row_green_osm, row_green_optimized] if row is not None]:
            for cell in row:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:val"), "clear")
                shading_elm.set(qn("w:color"), "auto")
                shading_elm.set(qn("w:fill"), "92D050")
                cell._tc.get_or_add_tcPr().append(shading_elm)

        # Inhalte je Tour
        for col_idx, tour_id in enumerate(tours_block, start=1):
            tour_data = tour_id_to_df[tour_id]
            tour_df = tour_data["tour_df"]

            # Kinder/Adresse füllen
            fornames = tour_df["fornames"].tolist()
            surnames = tour_df["surnames"].tolist()
            streets = tour_df["streets"].tolist()
            numbers = tour_df["housenumbers"].tolist()
            regions = tour_df["regions"].tolist()

            for i in range(min(9, len(fornames))):
                cell = table.rows[i + 2].cells[col_idx]
                if all(element[i] == "Platz ist frei!" for element in [fornames, surnames, streets, numbers]):
                    text = " \n \n "
                else:
                    text = f"{surnames[i]}, {fornames[i]}\n{streets[i]} {numbers[i]}\n{regions[i]}"

                # Clear existing text (important!)
                for p in cell.paragraphs:
                    p.clear() if hasattr(p, "clear") else None

                # Add text manually so we can control formatting
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.size = Pt(10)  # Adjust size here
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Or CENTER, if you prefer

            # KM-besetzt in grüne Zelle (Malteser)
            green_cell = table.rows[12].cells[col_idx]
            green_cell.text = str(tour_data["km_besetzt"])
            green_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # KM per tour for OSM
            if row_green_osm is not None and tour_id in google_distances:
                green_cell = row_green_osm[col_idx]
                green_cell.text = str(google_distances[tour_id])
                green_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # KM per tour for optimized
            if row_green_optimized is not None and tour_id in optimized_distances:
                green_cell = row_green_optimized[col_idx]
                green_cell.text = str(optimized_distances[tour_id])
                green_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Seitenumbruch falls weitere Touren folgen
        if end < num_tours:
            doc.add_page_break()

    # Save file in acceptable format
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer